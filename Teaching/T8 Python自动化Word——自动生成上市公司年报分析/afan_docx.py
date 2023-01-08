################################################
# 1 文字输出函数封装
def stock_revenue_status(stock_name, stock_id, stock_df):
    revenue = stock_df['营业收入'][-1]
    last_revenue = stock_df['营业收入'][-2]
    net_income = stock_df['净利润'][-1]
    last_income = stock_df['净利润'][-2]
    revenue_rate = revenue / last_revenue - 1
    net_income_rate = net_income / last_income - 1
    revenue_flag = '增长' if revenue_rate > 0 else '下降'
    net_income_flag = '增长' if net_income_rate > 0 else '下降'
    p1_text_1 = f'{stock_name}({stock_id})全年实现营业总收入{"%.2f" %(revenue/1e8)}亿元，同比{revenue_flag}{"%.2f" % (revenue_rate*100)}%；'
    p1_text_2 = f'实现净利润{"%.2f" %(net_income/1e8)}亿元，同比{net_income_flag}{"%.2f" % (net_income_rate*100)}%。'
    p1_text = p1_text_1 + p1_text_2
    return p1_text

def stock_efficiency_status(stock_name, stock_id, stock_df):
    net_rate = stock_df['利润率'][-1]
    last_net_rate = stock_df['利润率'][-2]
    net_rate_inc = net_rate - last_net_rate

    gross_rate = stock_df['毛利率'][-1]
    last_gross_rate = stock_df['毛利率'][-2]
    gross_rate_inc = gross_rate - last_gross_rate

    sale_rate = stock_df['销售费用率'][-1]
    last_sale_rate = stock_df['销售费用率'][-2]
    sale_rate_inc = sale_rate - last_sale_rate

    admin_rate = stock_df['管理费用率'][-1]
    last_admin_rate = stock_df['管理费用率'][-2]
    admin_rate_inc = admin_rate - last_admin_rate

    finance_rate = stock_df['财务费用率'][-1]
    last_finance_rate = stock_df['财务费用率'][-2]
    finance_rate_inc = finance_rate - last_finance_rate

    net_rate_flag = '增长' if net_rate_inc > 0 else '下降'
    gross_rate_flag = '增长' if gross_rate_inc > 0 else '下降'
    sale_rate_flag = '增长' if sale_rate_inc > 0 else '下降'
    admin_rate_flag = '增长' if admin_rate_inc > 0 else '下降'
    finance_rate_flag = '增长' if finance_rate_inc > 0 else '下降'

    p2_text_1 = f'净利率同期{net_rate_flag}{"%.2f" % (net_rate_inc*100)}pct至{"%.2f" % (net_rate*100)}%，'
    p2_text_2 = f'毛利率同期{gross_rate_flag}{"%.2f" % (gross_rate_inc*100)}pct至{"%.2f" % (gross_rate*100)}%。'
    p2_text_3 = f'费用率方面，销售费用率同期{sale_rate_flag}{"%.2f" % (sale_rate_inc*100)}pct至{"%.2f" % (sale_rate*100)}%，'
    p2_text_4 = f'管理费用率同期{admin_rate_flag}{"%.2f" % (admin_rate_inc*100)}pct至{"%.2f" % (admin_rate*100)}%，'
    p2_text_5 = f'财务费用率同期{finance_rate_flag}{"%.2f" % (finance_rate_inc*100)}pct至{"%.2f" % (finance_rate*100)}%。'
    p2_text = p2_text_1 + p2_text_2 + p2_text_3 + p2_text_4 + p2_text_5
    return p2_text

def stock_finance_status(stock_name, stock_id, stock_df):
    roe = stock_df['净资产收益率'][-1]
    last_roe = stock_df['净资产收益率'][-2]
    roe_inc = roe - last_roe

    net_rate = stock_df['利润率'][-1]
    last_net_rate = stock_df['利润率'][-2]
    net_rate_inc = net_rate - last_net_rate

    asset_turnover_rate = stock_df['资产周转率'][-1]
    last_asset_turnover_rate = stock_df['资产周转率'][-2]
    asset_turnover_rate_inc = asset_turnover_rate - last_asset_turnover_rate

    debt_asset_rate = stock_df['资产负债率'][-1]
    last_debt_asset_rate = stock_df['资产负债率'][-2]
    debt_asset_rate_inc = debt_asset_rate - last_debt_asset_rate

    roe_flag = '增长' if roe_inc > 0 else '下降'
    net_rate_flag = '增长' if net_rate_inc > 0 else '下降'
    asset_turnover_rate_flag = '增长' if asset_turnover_rate_inc > 0 else '下降'
    debt_asset_rate_flag = '增长' if debt_asset_rate_inc > 0 else '下降'

    p3_text_1 = f'ROE同期{roe_flag}{"%.2f" % (roe_inc*100)}pct至{"%.2f" % (roe*100)}%，根据杜邦分析法拆解，'
    p3_text_2 = f'净利率同期{net_rate_flag}{"%.2f" % (net_rate_inc*100)}pct至{"%.2f" % (net_rate*100)}%。'
    p3_text_3 = f'资产周转率同期{asset_turnover_rate_flag}{"%.2f" % (asset_turnover_rate_inc*100)}pct至{"%.2f" % (asset_turnover_rate*100)}%，'
    p3_text_4 = f'资产负债率同期{debt_asset_rate_flag}{"%.2f" % (debt_asset_rate_inc*100)}pct至{"%.2f" % (debt_asset_rate*100)}%。'

    p3_text = p3_text_1 + p3_text_2 + p3_text_3 + p3_text_4
    return p3_text

def stock_industry_status(stock_name, stock_id, industry_df, rank_index_list, index_direction_list, rank_threshold):
    # 先赋值，再更新
    p4_text = ''
    for index_name, direction in zip(rank_index_list, index_direction_list):
        # 注意这里方向是反过来的
        tmp_ascending = False if direction else True
        tmp_df = industry_df.sort_values(by=index_name, ascending=tmp_ascending)
        tmp_rank = tmp_df['股票简称'].tolist().index(stock_name)
        if tmp_rank + 1 <= rank_threshold:
            tmp_value = tmp_df.loc[stock_id, index_name]
            tmp_avg_value = tmp_df[index_name].mean()
            direction_flag = '远超过' if direction else '远低于'
            p4_text = f'公司的{index_name}表现良好，达到{"%.2f" % (tmp_value*100)}%，在行业内排名第{tmp_rank+1}，{direction_flag}行业平均的{"%.2f" % (tmp_avg_value*100)}%。'
            break
    return p4_text

################################################
# 2 表格输出函数封装
def get_finance_table(stock_df):
    import pandas as pd
    revenue_value = stock_df['营业收入'][1:]
    revenue_rate = stock_df['营业收入'].pct_change(1)[1:]
    EBIT_value = stock_df['息税前利润'][1:]
    EBIT_rate = stock_df['息税前利润'].pct_change(1)[1:]
    net_income_value = stock_df['净利润'][1:]
    net_income_rate = stock_df['净利润'].pct_change(1)[1:]
    finance_table = pd.concat([revenue_value, revenue_rate, EBIT_value, EBIT_rate, net_income_value, net_income_rate], axis=1).T
    finance_table.columns = finance_table.columns.str.slice(0, 4) + 'A'
    finance_table.index = ['营业收入', '(+/-)%', '经营利润(EBIT)', '(+/-)%', '净利润', '(+/-)%']
    return finance_table

def get_dupont_table(stock_df):
    import pandas as pd
    # 尽管不用算增速可以多1条数据，但是也要和其他数据对齐
    roe = stock_df['净资产收益率'][1:]
    net_rate = stock_df['利润率'][1:]
    asset_turnover_rate = stock_df['资产周转率'][1:]
    debt_asset_rate = stock_df['资产负债率'][1:]
    dupont_table = pd.concat([roe, net_rate, asset_turnover_rate, debt_asset_rate], axis=1).T
    dupont_table.columns = dupont_table.columns.str.slice(0, 4) + 'A'
    dupont_table.index = ['净资产收益率(ROE)', '利润率', '资产周转率', '资产负债率']
    return dupont_table

def add_df_document(df, doc, table_name):
    import pandas as pd
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1]+1, style='Light Shading Accent 1')
    # 增加表头
    hdr_cells = table.rows[0].cells
    # 第一行的第一个格子填充表名
    hdr_cells[0].text = table_name
    for col_idx, col_name in enumerate(df.columns):
        hdr_cells[col_idx+1].text = col_name
    # 增加内容，先按行循环
    for idx_idx, idx_name in enumerate(df.index):
        # 增加一行
        row_cells = table.rows[idx_idx+1].cells
        # 给这行填充列的值，先填充这行的指标名称
        row_cells[0].text = idx_name
        for col_idx, col_name in enumerate(df.columns):
            tmp_value = df.iloc[idx_idx, col_idx]
            # 如果指标名称有比率的情况
            if ('%' in idx_name) or ('率' in idx_name):
                row_cells[col_idx+1].text = "%.2f" % (tmp_value*100)
            else:  # 否则就简单按百万计算
                row_cells[col_idx+1].text = str(int(tmp_value / 1e6))
    return doc

################################################
# 3 word样式封装
def change_font(h, font_name=u'微软雅黑'):
    from docx.oxml.ns import qn
    h.style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    h.style.element.rPr.rFonts.set(qn("w:eastAsiaTheme"), font_name)

################################################
# 4 图片输出函数封装
def save_revenue_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    from mpl_toolkits.axes_grid1 import host_subplot
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    revenue_value = stock_df['营业收入'][1:]
    revenue_rate = stock_df['营业收入'].pct_change(1)[1:]
    x = revenue_value.index.str.slice(0, 4)
    y1 = revenue_value.values / 1e8
    y2 = revenue_rate.values * 100
    # 画图
    host = host_subplot(111)
    par1 = host.twinx()
    host.bar(x,y1,color='C0', label='营业收入(亿元)')
    par1.plot(x,y2, color='C1', label='营业收入yoy（%）', linewidth=3)
    host.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

def save_net_income_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    from mpl_toolkits.axes_grid1 import host_subplot
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    net_income_value = stock_df['净利润'][1:]
    net_income_rate = stock_df['净利润'].pct_change(1)[1:]
    x = net_income_value.index.str.slice(0, 4)
    y1 = net_income_value.values / 1e8
    y2 = net_income_rate.values * 100
    # 画图
    host = host_subplot(111)
    par1 = host.twinx()
    host.bar(x,y1,color='C0', label='净利润(亿元)')
    par1.plot(x,y2, color='C1', label='净利润yoy（%）', linewidth=3)
    host.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

def save_profit_rate_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    net_rate = stock_df['利润率'][1:]
    gross_rate = stock_df['毛利率'][1:]
    # 画图
    x = net_rate.index.str.slice(0, 4)
    y0 = net_rate.values * 100
    y1 = gross_rate.values * 100
    fig = plt.figure()
    plt.plot(x, y0, label='利润率（%）')
    plt.plot(x, y1, label='毛利率（%）')
    plt.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

def save_cost_rate_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    sale_rate = stock_df['销售费用率'][1:]
    admin_rate = stock_df['管理费用率'][1:]
    finance_rate = stock_df['财务费用率'][1:]
    # 画图
    x = sale_rate.index.str.slice(0, 4)
    y2 = sale_rate.values * 100
    y3 = admin_rate.values * 100
    y4 = finance_rate.values * 100
    fig = plt.figure()
    plt.plot(x, y2, label='销售费用率（%）')
    plt.plot(x, y3, label='管理费用率（%）')
    plt.plot(x, y4, label='财务费用率（%）')
    plt.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

def save_dupont1_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    roe = stock_df['净资产收益率'][1:]
    asset_turnover_rate = stock_df['资产周转率'][1:]
    net_rate = stock_df['利润率'][1:]
    debt_asset_rate = stock_df['资产负债率'][1:]
    # 画图
    x = roe.index.str.slice(0, 4)
    y0 = roe.values * 100
    y1 = asset_turnover_rate.values * 100
    y2 = net_rate.values * 100
    fig = plt.figure()
    plt.plot(x, y0, label='ROE（%）')
    plt.plot(x, y1, label='资产周转率（%）')
    plt.plot(x, y2, label='利润率（%）')
    plt.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

def save_dupont2_pic(stock_df, pic_name):
    # 导包
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif']=['SimHei'] #用来正常显示中文标签
    plt.rcParams['axes.unicode_minus']=False #用来正常显示负号
    # 计算
    roe = stock_df['净资产收益率'][1:]
    asset_turnover_rate = stock_df['资产周转率'][1:]
    net_rate = stock_df['利润率'][1:]
    debt_asset_rate = stock_df['资产负债率'][1:]
    # 画图
    x = roe.index.str.slice(0, 4)
    y2 = debt_asset_rate.values * 100
    fig = plt.figure()
    plt.plot(x, y2, label='资产负债率（%）')
    plt.legend(loc="upper left")
    plt.savefig(pic_name)
    # 清空画布
    plt.show()

################################################
# 5 word导出函数
def output_report(stock_id, report_year, total_df):
    ################################################
    # 设定参数
    report_date = f'{report_year}-12-31 00:00:00'
    stock_df = total_df.swaplevel().loc[stock_id].sort_index()
    stock_df = stock_df.loc[:report_date, :]
    stock_industry = stock_df['行业名称'][0]
    stock_data_row = stock_df.shape[0]
    # 如果只有1行不够！
    if stock_data_row > 1:
        stock_name = stock_df['股票简称'][0]
    else:  # raise是主动抛出异常
        raise Exception('不存在该股票的数据')
    # 行业方面设定重要性指标排序
    rank_index_list = ['净资产收益率', '利润率', '资产周转率', '资产负债率']
    # 1表示越大越好， 0表示越小越好
    index_direction_list = [1, 1, 1, 0]
    # 设定排名前几输出
    rank_threshold = 3
    # 行业数据
    report_year_df = total_df.loc[report_date]
    industry_df = report_year_df[report_year_df['行业名称']==stock_industry]
    
    ################################################
    # 生成图片
    import os
    root_path = '年报图表'
    stock_doc_path = f'{stock_name}-{stock_id}-{report_year}'
    pic_path = os.path.join(root_path, stock_doc_path)
    if not os.path.exists(pic_path):
        os.makedirs(pic_path)
    save_revenue_pic(stock_df, os.path.join(pic_path, '1.jpg'))
    save_net_income_pic(stock_df, os.path.join(pic_path, '2.jpg'))
    save_profit_rate_pic(stock_df, os.path.join(pic_path, '3.jpg'))
    save_cost_rate_pic(stock_df, os.path.join(pic_path, '4.jpg'))
    save_dupont1_pic(stock_df, os.path.join(pic_path, '5.jpg'))
    save_dupont2_pic(stock_df, os.path.join(pic_path, '6.jpg'))

    ################################################
    # 编辑导出
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.shared import Pt 

    document = Document()

    # 修改字体
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    # 标题部分
    heading = document.add_heading(f'{stock_name}({stock_id}){report_year}年报分析', 0)
    change_font(heading)

    # 第一段文字
    heading1 = document.add_heading('营收利润情况', level=1)
    change_font(heading1)
    p1_text = stock_revenue_status(stock_name, stock_id, stock_df)
    document.add_paragraph(p1_text)
    # 第一组图片
    para_pic1 = document.add_paragraph()
    para_pic1.add_run().add_picture(os.path.join(pic_path, '1.jpg'), width=Cm(7.5))
    para_pic1.add_run().add_picture(os.path.join(pic_path, '2.jpg'), width=Cm(7.5))

    # 第二段文字
    heading2 = document.add_heading('运营效率情况', level=1)
    change_font(heading2)
    p2_text = stock_efficiency_status(stock_name, stock_id, stock_df)
    document.add_paragraph(p2_text)
    # 第二组图片
    para_pic1 = document.add_paragraph()
    para_pic1.add_run().add_picture(os.path.join(pic_path, '3.jpg'), width=Cm(7.5))
    para_pic1.add_run().add_picture(os.path.join(pic_path, '4.jpg'), width=Cm(7.5))

    # 第三段文字
    heading3 = document.add_heading('杜邦分析', level=1)
    change_font(heading3)
    p3_text = stock_finance_status(stock_name, stock_id, stock_df)
    document.add_paragraph(p3_text)
    # 第三组图片
    para_pic1 = document.add_paragraph()
    para_pic1.add_run().add_picture(os.path.join(pic_path, '5.jpg'), width=Cm(7.5))
    para_pic1.add_run().add_picture(os.path.join(pic_path, '6.jpg'), width=Cm(7.5))

    # 第四段文字
    p4_text = stock_industry_status(stock_name, stock_id, industry_df, rank_index_list, index_direction_list, rank_threshold)
    if p4_text:  # 字符串不为''就是True
        heading4 = document.add_heading('行业情况', level=1)
        change_font(heading4)
        document.add_paragraph(p4_text)

    # 翻页获取表格
    # document.add_page_break()
    finance_table = get_finance_table(stock_df)
    dupont_table = get_dupont_table(stock_df)
    document = add_df_document(finance_table, document, '财务摘要（百万元）')
    document.add_paragraph('\n')
    document = add_df_document(dupont_table, document, '杜邦分析指标')

    # 保存
    document.save(stock_doc_path+ '年报分析.docx')